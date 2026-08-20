#Raspberry PI Portable Software Terminal Emulator by Jenierre Domingo
import subprocess
import os
from datetime import datetime
import serial
import time
import tkinter as tk
from tkinter import messagebox
from docx import Document
import re

#this is the directory where the script is located, and where the reports will be saved
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
REPORTS_DIR = os.path.join(SCRIPT_DIR, "reports")

#PORT = "/dev/ttyUSB0
BAUD = 9600 #Temporary loopback for GPIO Serial Ports

#these are all the units that can be selected, along with their serial settings
UNITS = {"2162": {"baud": 9600, "data_bits": 8, "parity": "N", "stop_bits": 1, "flow": "N"},
	 "2455": {"baud": 115200, "data_bits": 8, "parity": "N", "stop_bits": 1, "flow": "N"},
	 "2509": {"baud": 115200, "data_bits": 8, "parity": "N", "stop_bits": 1, "flow": "N"},
	 "2561": {"baud": 57600, "data_bits": 8, "parity": "N", "stop_bits": 1, "flow": "N"}, }

current_unit = None
putty_process = None
USE_USB_MODE = True
PORT = "/dev/ttyUSB0"
USB_PORT = "/dev/ttyUSB0"

#this function performs a loopback test on the serial port to check if it is working
def serial_loopback_detected():
	try:
		ser = serial.Serial(PORT, baudrate=BAUD, timeout=1)
		time.sleep(0.2)

		test_msg = b"PING\n"
		ser.write(test_msg)
		time.sleep(0.2)

		received = ser.readline()
		ser.close()

		return received == test_msg

	except Exception as e:
		print("Serial error:", e)
		return False

#this function clears the GUI screen by destroying all widgets
def clear_screen():
	for widget in root.winfo_children():
		widget.destroy()

#this function checks if the serial connection is active, either by checking for the USB port or by performing a loopback test
def serial_connection_active():
	if USE_USB_MODE:
		return os.path.exists(USB_PORT)

	return serial_loopback_detected()

serial_was_connected = serial_connection_active()
disconnect_popup_shown = False

#this function monitors the serial connection and shows a popup if it is disconnected, or if it is reconnected
def monitor_serial_connection():
	global serial_was_connected, putty_process
	#GPIO LOOPBACK MODE, Don't test the serial port while PuTTY owns it
	if not USE_USB_MODE:
		if putty_process is not None and putty_process.poll() is None:
			root.after(3000, monitor_serial_connection)
			return

	connected = serial_connection_active()

	if not connected and serial_was_connected:
		serial_was_connected = False

		if putty_process is not None and putty_process.poll() is None:
			putty_process.terminate()
			putty_process = None

		root.lift()
		root.attributes("-topmost", True)
		root.update()
		messagebox.showwarning("Serial Disconnected", "Serial disconnected.\n\nReconnect serial cable.")
		root.attributes("-topmost", False)
		show_no_connection_screen()

	elif connected and not serial_was_connected:
		serial_was_connected = True
		root.lift()
		root.attributes("-topmost", True)
		root.update()

		messagebox.showinfo("Serial Connection Detected", f"Serial connection detected on {PORT}.")
		root.attributes("-topmost", False)
		show_unit_buttons()

	root.after(3000, monitor_serial_connection)

#this allows the user to choose the unit type, and then displays the settings for that unit, along with a button to start the DSP session
def choose_unit(unit):
	global current_unit
	current_unit = unit
	settings = UNITS[unit]

	clear_screen()

	tk.Label(root, text=f"Selected Unit: {unit}", font=("Arial", 26)).pack(pady=30)

	tk.Label(root, text=f"Port: {PORT}\nBaud: {settings['baud']}", font=("Arial", 18)).pack(pady=20)

	tk.Button(root, text="Start DSP Session", font=("Arial", 22), width=18, height=2, command=start_session).pack(pady=15)

	tk.Button(root, text="Switch Unit", font=("Arial", 22), width=18, height=2, command=show_unit_buttons).pack(pady=15)

	tk.Checkbutton(root, text="Save DSP report to docx file", variable=save_log_var, font=("Arial",18)).pack(pady=10)

	tk.Label(root, text="Report File Name:", font=("Arial",16)).pack(pady=5)

	tk.Entry(root, textvariable=filename_var, font=("Arial",18), width=30).pack(pady=10)

	tk.Label(root, text="Unit Serial Number:", font=("Arial",18)).pack(pady=5)

	tk.Entry(root, textvariable=serial_number_var, font=("Arial",18), width=30).pack(pady=10)

	filename_var.set(f"{unit}_DSP_Report")

#if there is no serial connection, this function shows a screen with a message and a retry button
def show_no_connection_screen():
	clear_screen()
	tk.Label(root, text="No serial connection detected", font=("Arial",24)).pack(pady=60)

	tk.Button(root, text="Retry", font=("Arial",22), width=15, height=2, command=retry_connection).pack(pady=20)

#this prompts the user to retry the serial connection, and shows a message if it is successful or not
def retry_connection():
	global serial_was_connected

	if serial_connection_active():
		serial_was_connected = True

		messagebox.showinfo("Serial Connection Detected", f"Serial connection detected on {PORT}.")
		show_unit_buttons()
	else:
		serial_was_connected = False
		messagebox.showwarning("No Serial Connection", "No serial connection detected.\n\nPlease check the cable and try again.")
		show_no_connection_screen()

#this displays the unit selection buttons for the user to choose which unit they want to test
def show_unit_buttons():
	clear_screen()

	tk.Label(root, text="Select Unit", font=("Arial", 24)).pack(pady=30)

	for unit in UNITS:
		tk.Button(root, text=unit, font=("Arial", 22), width=15, height=2, command=lambda u=unit: choose_unit(u)).pack(pady=8)

#this function starts the DSP session by launching PuTTY with the selected unit's settings, and optionally saves the output to a Word document
def start_session():
	if current_unit is None:
		messagebox.showerror("Error", "No unit selected.")
		return

	settings = UNITS[current_unit]
	#Prompts DSP Session Process for USer
	messagebox.showinfo("Starting Session", f"Starting DSP session for {current_unit}\n\n"
		f"Unit: {current_unit}\n"
		f"Port: {PORT}\n"
		f"Baud Rate: {settings['baud']}\n"
		f"Data Bits: {settings['data_bits']}\n"
		f"Parity: {settings['parity']}\n"
		f"Stop Bits: {settings['stop_bits']}\n\n"

		f"Launching PuTTY...")

	print(f"Starting DSP session for {current_unit}")
	print(f"Using {PORT} at {UNITS[current_unit]['baud']} baud")

	serial_config = (
		f"{settings['baud']},"
		f"{settings['data_bits']},"
		f"{settings['parity']},"
		f"{settings['stop_bits']},"
		f"{settings['flow']}"
	)

	putty_command = [
		"putty",
		"-serial", PORT,
		"-sercfg", serial_config
	]

	if save_log_var.get():
		os.makedirs(REPORTS_DIR, exist_ok=True)
		user_filename = filename_var.get().strip()
		if user_filename == "":
			user_filename = f"{current_unit}_DSP_Report"
		if not user_filename.lower().endswith(".docx"):
			user_filename += ".docx"

		docx_file = os.path.join(REPORTS_DIR, user_filename)
		temp_log = "/tmp/dsp_temp.log"
		putty_command.extend(["-sessionlog", temp_log])

		print("REPORT LOGGING ENABLED")
		print("Temporary log:", temp_log)
		print("Final DOCX:", docx_file)

	print(f"Launching PuTTY for {current_unit}")
	print(" ".join(putty_command))

	try:
		global putty_process
		putty_process = subprocess.Popen(putty_command)
		putty_process.wait()
		if save_log_var.get():
			serial_number = serial_number_var.get().strip()
			print("\n========== REPORT DEBUG ==========")
			print("PuTTY closed")
			print("Current working directory:", os.getcwd())
			print("Reports directory:", REPORTS_DIR)
			print("Temp log:", temp_log)
			print("Temp log exists:", os.path.exists(temp_log))
			print("DOCX  path:", docx_file)
			print("====================\n")

			if not os.path.exists(temp_log):
				messagebox.showerror("Report Error", "PuTTY did not create the temporary DSP log.\n\n" f"Expected location:\n{temp_log}")
				return

			convert_log_to_docx(temp_log, docx_file, current_unit, serial_number)
			if os.path.exists(docx_file):
				print("SUCCESS:DOCX exists")
				if os.path.exists(temp_log):
					os.remove(temp_log)
				messagebox.showinfo("Report Saved", f"Word report saved as:\n\n{docx_file}")
			else:
				messagebox.showerror("Report Error", "The Word document was not created.")
		root.after(1000, maximize_putty)
	except Exception as e:
		messagebox.showerror("PuTTY Launch Error", str(e))

#this fullscreens putty after it is launched, but keeps the windows bar visible
def maximize_putty():
	subprocess.run([
		"wmctrl", "-r", "PuTTY", "-b", "add,maximized_vert,maximized_horz"])

#this function converts the temporary DSP log text file to a Word document, and adds the unit, serial number, and date to the report
def convert_log_to_docx(txt_path, docx_path, unit, serial_number):
	document = Document()

	document.add_heading("DSP Test Report", level=1)

	document.add_paragraph(f"Unit: {unit}")
	document.add_paragraph(f"Serial Number: {serial_number}")
	document.add_paragraph(f"Date: {datetime.now().strftime('%m/%d/%Y %I:%M %p')}")

	document.add_heading("DSP Output", level=2)

	with open(txt_path, "r", errors="replace") as file:
		dsp_data = file.read()
	dsp_data = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', dsp_data)

	document.add_paragraph(dsp_data)
	document.save(docx_path)
	print("DOCX SAVE FINISHED")
	print("DOCX exists:", os.path.exists(docx_path))

root = tk.Tk()
root.tk.call('tk', 'scaling', 1.75)
root.title("Portable DSP Emulator")
root.geometry("800x480")
#Maximize while keeping windows bar
try:
	root.attributes("-zoomed", True)
except tk.TclError:
	root.state("zoomed")


save_log_var = tk.BooleanVar(value=True)
filename_var = tk.StringVar(value="")
serial_number_var = tk.StringVar(value="")

print("Checking serial connection on /dev/serial0...")

if serial_connection_active():
	serial_was_connected = True
	root.lift()
	root.attributes("-topmost",True)
	root.update()
	messagebox.showinfo("Serial Connection Detected", f"Serial connection detected on {PORT}.")
	root.attributes("-topmost", False)
	show_unit_buttons()
else:
	serial_was_connected = False
	show_no_connection_screen()

root.after(3000, monitor_serial_connection)
root.mainloop()
